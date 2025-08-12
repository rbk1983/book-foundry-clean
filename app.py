# app.py ‚Äî Book Auto-Generator (tabs + autosave/resume + post-process)
import os, sys, time, math, hashlib, base64, json, re, io, random, datetime as _dt
from typing import List, Dict, Any, Optional, Tuple

import streamlit as st

# ========== Page & watchdog ==========
st.set_page_config(page_title="Book Auto-Generator", layout="wide")

# ---- Emergency unlock support ----
if "busy" not in st.session_state:
    st.session_state.busy = False
if "last_tick" not in st.session_state:
    st.session_state.last_tick = time.time()
now = time.time()
if st.session_state.busy and (now - st.session_state.last_tick) > 180:
    st.warning("Previous run appears stalled; unlocking UI.")
    st.session_state.busy = False

# ========== Diagnostics (sidebar) ==========
with st.sidebar:
    st.markdown("### üîç Diagnostics")
    st.write("Python:", sys.version.split()[0])
    try:
        import openai as _oa
        st.write("openai pkg:", getattr(_oa, "__version__", "unknown"))
    except Exception as e:
        st.write("openai import error:", e)
    try:
        import httpx as _hx
        st.write("httpx:", getattr(_hx, "__version__", "unknown"))
    except Exception as e:
        st.write("httpx import error:", e)
    has_tavily = False
    try:
        has_tavily = bool(st.secrets.get("TAVILY_API_KEY", ""))
    except Exception:
        has_tavily = bool(os.getenv("TAVILY_API_KEY", ""))
    st.write("Tavily key:", "‚úÖ" if has_tavily else "‚Äî")
    if st.button("‚ö° Force unlock UI"):
        st.session_state.busy = False
        st.rerun()

os.makedirs("data", exist_ok=True)

# ========== Secrets helper ==========
def _sec(name: str) -> Optional[str]:
    try:
        if name in st.secrets:
            return st.secrets[name]
    except Exception:
        pass
    return os.getenv(name)

# ========== OpenAI client + robust retry ==========
from openai import OpenAI
OPENAI_KEY = _sec("OPENAI_API_KEY")
if not OPENAI_KEY:
    st.error("Missing OPENAI_API_KEY in Streamlit secrets.")
    st.stop()
client = OpenAI(api_key=OPENAI_KEY)

def _backoff_sleep(attempt: int):
    base = min(8, 0.5 * (2 ** attempt))
    time.sleep(base + random.uniform(0, 0.5))

def chat_with_retry(messages, model, temperature, max_tokens, tries: int = 7) -> str:
    last_err = None
    for attempt in range(tries):
        try:
            resp = client.chat.completions.create(
                model=model, temperature=temperature,
                messages=messages, max_tokens=max_tokens
            )
            return resp.choices[0].message.content
        except Exception as e:
            last_err = e
            _backoff_sleep(attempt)
    raise RuntimeError(f"OpenAI chat failed after {tries} attempts: {last_err}")

def embed_with_retry(texts: List[str], model: str, tries: int = 7) -> List[List[float]]:
    last_err = None
    for attempt in range(tries):
        try:
            resp = client.embeddings.create(model=model, input=texts)
            return [d.embedding for d in resp.data]
        except Exception as e:
            last_err = e
            _backoff_sleep(attempt)
    raise RuntimeError(f"OpenAI embeddings failed after {tries} attempts: {last_err}")

# ========== GitHub helpers ==========
import httpx

def gh_headers():
    tok = _sec("GITHUB_TOKEN")
    return {
        "Authorization": f"Bearer {tok}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
    }

def gh_repo_info():
    owner  = _sec("GH_REPO_OWNER")
    repo   = _sec("GH_REPO_NAME")
    branch = _sec("GH_BRANCH")
    return owner, repo, branch

def gh_put_file(path_rel: str, content_bytes: bytes, message: str) -> dict:
    owner, repo, branch = gh_repo_info()
    if not (owner and repo):
        raise RuntimeError("GH_REPO_OWNER or GH_REPO_NAME missing in secrets.")
    url = f"https://api.github.com/repos/{owner}/{repo}/contents/{path_rel}"
    params = {"ref": branch} if branch else None
    with httpx.Client(timeout=90.0) as c:
        r0 = c.get(url, params=params, headers=gh_headers())
        sha = r0.json().get("sha") if r0.status_code == 200 else None
        if r0.status_code not in (200, 404):
            r0.raise_for_status()
        payload = {
            "message": message,
            "content": base64.b64encode(content_bytes).decode("utf-8"),
        }
        if branch:
            payload["branch"] = branch
        if sha:
            payload["sha"] = sha
        r = c.put(url, headers=gh_headers(), data=json.dumps(payload))
        r.raise_for_status()
        return r.json()

def gh_list_any(path_rel: str) -> List[dict]:
    owner, repo, branch = gh_repo_info()
    if not (owner and repo):
        return []
    url = f"https://api.github.com/repos/{owner}/{repo}/contents/{path_rel}"
    params = {"ref": branch} if branch else None
    with httpx.Client(timeout=60.0) as c:
        r = c.get(url, params=params, headers=gh_headers())
        if r.status_code == 404:
            return []
        r.raise_for_status()
        return r.json()

def gh_get_file(path_rel: str) -> bytes:
    owner, repo, branch = gh_repo_info()
    if not (owner and repo):
        raise RuntimeError("Missing GH repo settings")
    url = f"https://api.github.com/repos/{owner}/{repo}/contents/{path_rel}"
    params = {"ref": branch} if branch else None
    with httpx.Client(timeout=90.0) as c:
        r = c.get(url, params=params, headers=gh_headers())
        r.raise_for_status()
        data = r.json()
        if data.get("encoding") == "base64" and "content" in data:
            return base64.b64decode(data["content"])
        raise RuntimeError("Unexpected GitHub get-file response format")

# ========== Text loaders & chunking ==========
from pypdf import PdfReader
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from markdown import markdown
from bs4 import BeautifulSoup
from langchain_text_splitters import RecursiveCharacterTextSplitter

def load_text_from_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            reader = PdfReader(path)
            parts = []
            for p in reader.pages:
                try:
                    parts.append(p.extract_text() or "")
                except Exception:
                    parts.append("")
            return "\n".join(parts).strip()
        elif ext == ".docx":
            doc = DocxDocument(path)
            return "\n".join(p.text for p in doc.paragraphs).strip()
        elif ext in (".md",".markdown"):
            html = markdown(open(path, "r", encoding="utf-8", errors="ignore").read())
            return BeautifulSoup(html, "html.parser").get_text("\n").strip()
        else:
            return open(path, "r", encoding="utf-8", errors="ignore").read().strip()
    except Exception as e:
        raise RuntimeError(f"Failed to read {os.path.basename(path)}: {e}")

def chunk_text(text: str, chunk_size=1000, overlap=150) -> List[str]:
    if not text or not text.strip():
        return []
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size, chunk_overlap=overlap,
        separators=["\n\n","\n",". ","? ","! "," ",""]
    )
    pieces = [t.strip() for t in splitter.split_text(text)]
    return [p for p in pieces if p]

# ========== Embeddings & retrieval ==========
def embed_texts(texts: List[str], model: str) -> List[List[float]]:
    texts = [t for t in texts if t and t.strip()]
    if not texts:
        return []
    return embed_with_retry(texts, model=model)

def cosine_sim(a: List[float], b: List[float]) -> float:
    dot = 0.0; na = 0.0; nb = 0.0
    for x, y in zip(a, b):
        dot += x * y
        na += x*x
        nb += y*y
    if na == 0 or nb == 0: return 0.0
    return dot / (math.sqrt(na)*math.sqrt(nb))

def sha16(s: str) -> str:
    return hashlib.sha256(s.encode("utf-8")).hexdigest()[:16]

# ========== Tavily search ==========
def tavily_search(query: str,
                  max_results: int = 6,
                  include_domains: Optional[List[str]] = None,
                  exclude_domains: Optional[List[str]] = None,
                  search_depth: str = "advanced",
                  days: Optional[int] = None,
                  tries: int = 6) -> List[Dict[str, Any]]:
    api_key = _sec("TAVILY_API_KEY")
    if not api_key:
        return []
    payload = {
        "api_key": api_key,
        "query": query,
        "max_results": max_results,
        "search_depth": search_depth,
        "include_answer": False,
        "include_raw_content": False,
        "source": "book-autogen"
    }
    if include_domains: payload["include_domains"] = include_domains
    if exclude_domains: payload["exclude_domains"] = exclude_domains
    if days is not None: payload["days"] = days

    last_err = None
    for attempt in range(tries):
        try:
            with httpx.Client(timeout=75.0) as c:
                r = c.post("https://api.tavily.com/search", json=payload)
                r.raise_for_status()
                data = r.json()
                results = data.get("results", [])
                out = []
                for rr in results:
                    out.append({
                        "url": rr.get("url"),
                        "title": rr.get("title"),
                        "content": rr.get("content", ""),
                        "score": rr.get("score", 0.0)
                    })
                return out
        except Exception as e:
            last_err = e
            _backoff_sleep(attempt)
    return []

# ========== Session state ==========
if "records" not in st.session_state:
    st.session_state.records = []  # [{id, text, tags, source, emb}]
if "model" not in st.session_state:
    st.session_state.model = "gpt-4o"
if "emb_model" not in st.session_state:
    st.session_state.emb_model = "text-embedding-3-large"

records: List[Dict[str,Any]] = st.session_state.records

# ========== Small utils ==========
def safe_title_slug(title: str) -> str:
    return re.sub(r"[^A-Za-z0-9_\-]+", "_", (title or "Book").strip()) or "Book"

def draft_dir(title: str, ts: int) -> str:
    return f"outputs/drafts/{safe_title_slug(title)}_{ts}"

def wc(t: str) -> int:
    return len((t or "").split())

# ========== Retrieval helpers ==========
def retrieve(query: str, k: int, tags: Optional[List[str]]=None) -> List[Dict[str,Any]]:
    if not records or not query.strip():
        return []
    qvecs = embed_texts([query], model=st.session_state.emb_model)
    if not qvecs:
        return []
    qvec = qvecs[0]
    scored = []
    tags_lower = set([t.lower() for t in (tags or [])])
    # source diversity: retain highest 1-2 per source
    bucket = {}
    for r in records:
        if tags_lower and not (set([t.lower() for t in r["tags"]]) & tags_lower):
            continue
        s = cosine_sim(qvec, r["emb"])
        src = r.get("source","")
        bucket.setdefault(src, [])
        bucket[src].append((s, r))
    picks = []
    for src, arr in bucket.items():
        arr.sort(key=lambda x: x[0], reverse=True)
        picks.extend(arr[:2])  # at most 2 per source
    picks.sort(key=lambda x: x[0], reverse=True)
    return [r for _,r in picks[:k]]

def make_context(hits: List[Dict[str,Any]]) -> str:
    out = []
    seen_people = set()
    for i,h in enumerate(hits):
        tagline = ",".join(h.get('tags',[]))
        src = h.get('source','')
        out.append(f"[S{i+1} | {tagline} | {src}] {h['text'][:1200]}")
        person_guess = re.findall(r"^[A-Z][a-z]+ [A-Z][a-z]+", h['text'][:100])
        if person_guess:
            seen_people.add(person_guess[0])
    return "\n\n".join(out) if out else "(no retrieved context)"

# ========== Outline & writing helpers ==========
def plan_outline(context: str, book_title: str, thesis: str, style: str, num_chapters: int, model: str) -> str:
    prompt = f"""
You are a senior acquisitions editor. Propose a complete outline (chapters + brief synopses) for a 250‚Äì300 page book.
Title: {book_title}
Thesis: {thesis}
Voice & Style: {style}
Desired total chapters: {num_chapters}

RULES
- The outline MUST include an **Introduction** as Chapter 1 and a **Final Conclusion** as the last chapter.
- Total chapters MUST equal {num_chapters}.
- Each chapter gets a short, actionable synopsis (2‚Äì4 sentences).
- Output format:
Chapter 1: Introduction ‚Äî 2‚Äì4 sentence synopsis
Chapter 2: Title ‚Äî synopsis
...
Chapter {num_chapters}: Final Conclusion ‚Äî synopsis

Context excerpts (may be empty):
{context}
""".strip()
    msgs = [
        {"role":"system","content":"You are an expert editor who creates commercially strong nonfiction outlines."},
        {"role":"user","content":prompt}
    ]
    return chat_with_retry(msgs, model=model, temperature=0.4, max_tokens=2000)

def parse_outline_lines(text: str) -> List[Dict[str,Any]]:
    chapters = []
    lines = text.replace("\r","\n").split("\n")
    current = None
    for line in lines:
        m = re.match(r"^\s*Chapter\s+(\d+)\s*:\s*(.+)$", line.strip(), re.IGNORECASE)
        mdh = re.match(r"^\s*#\s+(.+)$", line.strip())
        if m or mdh:
            if m:
                num = m.group(1).strip()
                tail = m.group(2).strip()
                title = tail
                synopsis = ""
                if "‚Äî" in tail:
                    parts = tail.split("‚Äî",1)
                    title = parts[0].strip()
                    synopsis = parts[1].strip()
                elif "-" in tail:
                    parts = tail.split("-",1)
                    title = parts[0].strip()
                    synopsis = parts[1].strip()
            else:
                num = str(len(chapters)+1)
                title = mdh.group(1).strip()
                synopsis = ""
            current = {"num":num, "title":title, "synopsis":synopsis, "sections":[]}
            chapters.append(current); continue
        bullet = re.match(r"^\s*[-*+]\s+(.+)$", line)
        subh   = re.match(r"^\s*###\s+(.+)$", line)
        enum   = re.match(r"^\s*\d+[\.)]\s+(.+)$", line)
        if current and (bullet or subh or enum):
            sec_title = (bullet or subh or enum).group(1).strip()
            if sec_title:
                current["sections"].append(sec_title)
    return chapters

def chapter_section_plan(ch_num: str, ch_title: str, ch_synopsis: str, target_words: int, max_sections: int, context: str, forced_sections: Optional[List[str]]=None, model: str="gpt-4o") -> str:
    if forced_sections:
        fs = [s.strip() for s in forced_sections if s.strip()]
        fs = [s for s in fs if s.lower() not in ("conclusion", "final conclusion")]
        if fs and fs[-1].lower() != "conclusion":
            fs.append("Conclusion")
        return "\n".join([f"{i+1}. {t}" for i,t in enumerate(fs[:max_sections])])

    prompt = f"""
Draft a concise section plan for Chapter {ch_num}: "{ch_title}" (~{target_words} words).
Synopsis: {ch_synopsis}
Create up to {max_sections} sections (ideally 3‚Äì4). Reserve the last section for "Conclusion".
Avoid section titles that are synonyms of Conclusion or Summary.

Output: numbered list only.
Context (may be empty):
{context}
""".strip()
    msgs = [
        {"role":"system","content":"You are an expert nonfiction book outliner."},
        {"role":"user","content":prompt}
    ]
    return chat_with_retry(msgs, model=model, temperature=0.35, max_tokens=500)

def draft_section(si: int, sections_total: int, ch_num: str, ch_title: str, ch_synopsis: str, section_words: int, section_plan: str, context: str, urls: List[str], guidance: str, recent_cutoff: int, model: str) -> str:
    urls_block = "\n".join(f"- {u}" for u in urls) if urls else "(none)"
    guidance_text = guidance.strip() if guidance else ""
    recency_line = ""
    if isinstance(recent_cutoff, int) and recent_cutoff > 0:
        recency_line = f"Prefer sources published in {recent_cutoff} or later when citing web material."

    prompt = f"""
Write Section {si} (~{section_words} words) for Chapter {ch_num}: "{ch_title}".
Follow this approved section plan:
{section_plan}

Synopsis: {ch_synopsis}

Context from my books (may be empty; paraphrase freely; you may quote verbatim only with attribution):
{context}

Web references (cite ONLY if you actually use a fact; add inline Markdown links): 
{urls_block}

Additional research guidance (apply if web links are provided):
{guidance_text}
{recency_line}

MANDATORY RULES
- Write ONLY Section {si}. Do not draft other sections.
- Keep a single "## Conclusion" for the final section ONLY. If this is not the last section, do NOT write any conclusion section.
- Include at least 1‚Äì2 **verbatim quotes** from my books where relevant, with attribution:
  ‚Äúquote‚Äù ‚Äî Full Name, Role/Title at Hotel/Restaurant (Country)
- Paraphrased book content: no citation.
- Web facts/figures: add a Markdown hyperlink at the point of use.
- Use crisp headings (## ... ; ### for subheads). Avoid more than 3‚Äì4 H2 sections in the chapter total.
- Maintain voice; avoid repetition; end with a segue to the next section.
""".strip()
    msgs = [
        {"role":"system","content":"You are an expert long-form writing assistant. Reply in Markdown."},
        {"role":"user","content":prompt}
    ]
    return chat_with_retry(msgs, model=model, temperature=0.6, max_tokens=1500)

def stitch_chapter(ch_num: str, ch_title: str, target_words: int, sections: List[str], model: str) -> str:
    prompt = f"""
Combine the sections below into a cohesive chapter (~{target_words} words).
Smooth transitions, remove duplicates, normalize headings.

CRITICAL RULES
- Exactly ONE "## Conclusion" at the end. If earlier sections added any conclusion-like headings, merge their content into regular sections and remove the extra heading.
- Preserve all inline attributions for verbatim quotes from my books:
  ‚Äúquote‚Äù ‚Äî Full Name, Role/Title at Hotel/Restaurant (Country)
- Keep web hyperlinks where used; do not invent links; do not add citations to paraphrased book content.

Sections:
{"\n\n---\n\n".join(sections)}
""".strip()
    msgs = [
        {"role":"system","content":"You are a meticulous editor. Reply in Markdown."},
        {"role":"user","content":prompt}
    ]
    return chat_with_retry(msgs, model=model, temperature=0.3, max_tokens=3200)

def dedup_boundary(old_text: str, new_text: str, lookback: int=160) -> str:
    old = (old_text or "").strip()
    newt = (new_text or "").strip()
    if not old: return newt
    old_end = old[-lookback:].lower()
    new_start = newt[:lookback].lower()
    if old_end in new_start or new_start in old_end:
        for i in range(min(lookback, len(newt))):
            if old.endswith(newt[:i]):
                return newt[i:]
    return newt

# ========== Markdown ‚Üí Docx ==========
from docx import Document as DocxDocument
def add_toc(paragraph):
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t'); t.text = "Table of Contents (open in Word > References > Update Table)"
    r.append(t); fld.append(r)
    paragraph._p.addnext(fld)

def md_to_docx(md_text: str, title: str, thesis: str, author: Optional[str] = None) -> bytes:
    doc = DocxDocument()
    h = doc.add_heading(title, level=0); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if author:
        a = doc.add_paragraph(author); a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = doc.add_paragraph(f"Thesis: {thesis}"); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    p_toc = doc.add_paragraph(); add_toc(p_toc); doc.add_page_break()
    lines = md_text.replace("\r","\n").split("\n")
    for line in lines:
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        else:
            para = doc.add_paragraph()
            last = 0
            for m in re.finditer(r"\[([^\]]+)\]\(([^)]+)\)", line):
                before = line[last:m.start()]
                if before: para.add_run(before)
                text = m.group(1); url = m.group(2)
                run = para.add_run(text); run.font.underline = True
                para.add_run(f" ({url})")
                last = m.end()
            tail = line[last:]
            if tail: para.add_run(tail)
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

# ========== Autosave / Resume ==========
def save_manifest(title: str, ts: int, manifest: Dict[str,Any]):
    path = f"{draft_dir(title, ts)}/manifest.json"
    gh_put_file(path, json.dumps(manifest, ensure_ascii=False, indent=2).encode("utf-8"), "Update manifest")

def save_chapter_draft(title: str, ts: int, ch_num: str, ch_title: str, text: str):
    name = f"ch_{int(ch_num):02d}_{safe_title_slug(ch_title)}.md"
    path = f"{draft_dir(title, ts)}/{name}"
    gh_put_file(path, text.encode("utf-8"), f"Add {name}")

def find_latest_manifest() -> Optional[Tuple[str,int,Dict[str,Any]]]:
    base = "outputs/drafts"
    items = gh_list_any(base)
    dirs = [i for i in items if i.get("type") == "dir"]
    latest = None; latest_ts = -1; latest_path = None
    for d in dirs:
        name = d["name"]
        m = re.search(r"_(\d+)$", name)
        if not m: continue
        ts = int(m.group(1))
        if ts > latest_ts:
            latest_ts = ts
            latest_path = d["path"]
    if latest_path:
        try:
            bytes_man = gh_get_file(f"{latest_path}/manifest.json")
            j = json.loads(bytes_man.decode("utf-8"))
            return latest_path, latest_ts, j
        except Exception:
            return None
    return None

# ========== Generator ==========
def build_outline_and_chapters(thesis_ctx: str, outline_mode: str, pasted_outline: str, outline_file,
                               book_title: str, thesis: str, style: str, num_chapters: int, model: str) -> List[Dict[str,Any]]:
    if outline_mode == "Auto-generate":
        st.info("Generating outline‚Ä¶")
        outline_text = plan_outline(thesis_ctx, book_title, thesis, style, num_chapters, model)
        st.markdown("#### Proposed Outline")
        st.code(outline_text)
        parsed_chapters = parse_outline_lines(outline_text)
    elif outline_mode == "Paste outline":
        outline_text = pasted_outline or ""
        st.markdown("#### Using your pasted outline")
        st.code(outline_text)
        parsed_chapters = parse_outline_lines(outline_text)
    elif outline_mode == "Upload outline file":
        if not outline_file:
            st.error("Please upload an outline file or switch modes.")
            st.stop()
        bytes_io = outline_file.read()
        tmp = os.path.join("data", f"outline_{int(time.time())}_{outline_file.name}")
        with open(tmp, "wb") as f:
            f.write(bytes_io)
        try:
            text = load_text_from_file(tmp)
        except Exception as e:
            st.error(f"Could not read outline: {e}")
            st.stop()
        outline_text = text
        st.markdown("#### Parsed outline source (raw)")
        st.code(outline_text[:4000] + ("..." if len(outline_text) > 4000 else ""))
        parsed_chapters = parse_outline_lines(outline_text)
    else:
        parsed_chapters = []

    def coerce_intro_conclusion(chaps: List[Dict[str,Any]], total: int) -> List[Dict[str,Any]]:
        arr = [c for c in chaps if c.get("title")]
        if not arr or "intro" not in arr[0]["title"].lower():
            arr.insert(0, {"num":"1","title":"Introduction","synopsis": arr[0]["synopsis"] if arr else "", "sections": []})
        if "conclusion" not in arr[-1]["title"].lower():
            arr.append({"num":str(len(arr)+1),"title":"Final Conclusion","synopsis":"","sections":[]})
        if len(arr) > total:
            middle_keep = total - 2
            arr = [arr[0]] + arr[1:1+middle_keep] + [arr[-1]]
        elif len(arr) < total:
            to_add = total - len(arr)
            for i in range(to_add):
                arr.insert(-1, {"num":"0","title":f"Chapter Placeholder {i+1}","synopsis":"","sections":[]})
        for i, c in enumerate(arr, start=1):
            c["num"] = str(i)
        return arr

    chapters = coerce_intro_conclusion(parsed_chapters, int(num_chapters))
    st.markdown("#### Outline Preview")
    for c in chapters:
        st.markdown(f"- **Chapter {c['num']}: {c['title']}** ‚Äî {c.get('synopsis','')}")
        secs = c.get("sections") or []
        if secs:
            for s in secs:
                st.markdown(f"  - {s}")
    if not chapters:
        st.error("No chapters parsed; please adjust your outline or settings.")
        st.stop()
    return chapters

def generate_book(run_mode: str,
                  model: str,
                  emb_model: str,
                  records: List[Dict[str,Any]],
                  book_title: str, thesis: str, style: str,
                  outline_mode: str, pasted_outline: str, outline_file,
                  num_chapters: int, words_per_chapter: int,
                  top_k: int, chunk_size: int, chunk_overlap: int,
                  use_tavily: bool, research_guidance: str, recent_year_cutoff: int,
                  include_domains_txt: str, exclude_domains_txt: str,
                  urls_list: List[str], max_sources_per_chapter: int,
                  temperature: float):
    st.session_state.busy = True
    try:
        st.session_state.last_tick = time.time()

        thesis_hits = retrieve(thesis or book_title, k=top_k, tags=None) if records else []
        thesis_ctx = make_context(thesis_hits)

        if run_mode == "fresh":
            chapters = build_outline_and_chapters(thesis_ctx, outline_mode, pasted_outline, outline_file,
                                                 book_title, thesis, style, num_chapters, model)
            ts = int(time.time())
            author = _sec("BOOK_AUTHOR") or ""
            manifest = {
                "title": book_title, "author": author,
                "thesis": thesis, "style": style,
                "num_chapters": int(num_chapters),
                "words_per_chapter": int(words_per_chapter),
                "chapters": chapters, "completed": [], "next_index": 0,
                "urls_mode": "manual+auto" if (use_tavily or urls_list) else "none",
                "use_tavily": bool(use_tavily),
                "research_guidance": research_guidance,
                "recent_year_cutoff": int(recent_year_cutoff),
                "include_domains": [d.strip() for d in include_domains_txt.split(",") if d.strip()],
                "exclude_domains": [d.strip() for d in exclude_domains_txt.split(",") if d.strip()],
                "top_k": int(top_k), "chunk_size": int(chunk_size), "chunk_overlap": int(chunk_overlap),
                "temperature": float(temperature), "timestamp": ts
            }
            gh_put_file(f"{draft_dir(book_title, ts)}/_header.txt", f"Draft start: {book_title} @ {ts}".encode("utf-8"), "Init draft folder")
            save_manifest(book_title, ts, manifest)
        else:
            found = find_latest_manifest()
            if not found:
                st.error("No prior manifest found in outputs/drafts/. Start a fresh run first.")
                return
            _, ts, manifest = found
            chapters = manifest["chapters"]

        progress = st.progress(0.0, text="Starting‚Ä¶")
        manuscript_parts = [f"# {manifest['title']}\n\n*Author:* {manifest.get('author','')}\n\n*Thesis:* {manifest.get('thesis','')}\n"]
        prior_text = ""
        include_domains = manifest.get("include_domains") or [d.strip() for d in include_domains_txt.split(",") if d.strip()]
        exclude_domains = manifest.get("exclude_domains") or [d.strip() for d in exclude_domains_txt.split(",") if d.strip()]
        start_idx = manifest.get("next_index", 0)
        total_ch = len(chapters)

        for idx in range(start_idx, total_ch):
            st.session_state.last_tick = time.time()
            c = chapters[idx]
            ch_num = c["num"]; ch_title = c["title"]; ch_syn = c.get("synopsis","")
            progress.progress(idx/total_ch, text=f"Drafting Chapter {ch_num}: {ch_title}")

            query = f"{manifest['title']} | {ch_title} | {ch_syn} | continuity: {prior_text[-600:] if prior_text else ''}"
            hits = retrieve(query, k=manifest.get("top_k", top_k), tags=["Books"]) if records else []
            ctx = make_context(hits)

            manual_urls = urls_list[:] if urls_list else []
            tavily_urls = []
            if manifest.get("use_tavily") and _sec("TAVILY_API_KEY"):
                rq_parts = [f"Topic: {ch_title}", f"Synopsis: {ch_syn}", f"Book thesis: {manifest.get('thesis','')}", f"Style: {manifest.get('style','')}"]
                if (manifest.get("research_guidance") or "").strip():
                    rq_parts.append(f"Guidance: {manifest['research_guidance'].strip()}")
                research_query = " | ".join(rq_parts)
                days_window = None
                ryc = manifest.get("recent_year_cutoff", 0)
                if isinstance(ryc, int) and ryc > 0:
                    start = _dt.datetime(ryc, 1, 1)
                    days_window = max(1, (_dt.datetime.utcnow() - start).days)
                raw = tavily_search(
                    query=research_query,
                    max_results=max_sources_per_chapter,
                    include_domains=include_domains,
                    exclude_domains=exclude_domains,
                    search_depth="advanced",
                    days=days_window
                )
                raw_sorted = sorted(raw, key=lambda r: -float(r.get("score",0.0)))
                tavily_urls = [r["url"] for r in raw_sorted if r.get("url")]

            effective_urls = []
            seen = set()
            for u in (manual_urls + tavily_urls):
                if u not in seen:
                    effective_urls.append(u); seen.add(u)

            forced_secs = c.get("sections") if isinstance(c.get("sections"), list) and c.get("sections") else None
            sections_target = max(3, min(4, int(round(manifest.get("words_per_chapter", 3000)/1000))))
            plan_text = chapter_section_plan(ch_num, ch_title, ch_syn, manifest.get("words_per_chapter", 3000), sections_target, ctx, forced_sections=forced_secs, model=model)

            sec_texts = []
            per_sec = max(700, min(1000, int(manifest.get("words_per_chapter", 3000)/sections_target)))
            total_sections_to_write = len(forced_secs) if forced_secs else sections_target
            for si in range(1, total_sections_to_write+1):
                s_txt = draft_section(si, total_sections_to_write, ch_num, ch_title, ch_syn, per_sec, plan_text, ctx, effective_urls, manifest.get("research_guidance",""), manifest.get("recent_year_cutoff",0), model=model)
                if sec_texts:
                    s_txt = dedup_boundary(sec_texts[-1], s_txt, lookback=160)
                sec_texts.append(s_txt)

            ch_text = stitch_chapter(ch_num, ch_title, manifest.get("words_per_chapter", 3000), sec_texts, model=model)

            passes=0
            while wc(ch_text) < int(manifest.get("words_per_chapter", 3000)*0.95) and passes < 4:
                cont_prompt = f"""Continue Chapter {ch_num}: "{ch_title}" seamlessly from where it stops.
Do NOT repeat prior text. Keep headings; add 1‚Äì2 subsections if appropriate.
Aim for ~{per_sec}‚Äì{per_sec+200} words. Current ending:
{ch_text[-900:]}"""
                msgs = [
                    {"role":"system","content":"You are an expert long-form writing assistant. Reply in Markdown."},
                    {"role":"user","content":cont_prompt}
                ]
                cont = chat_with_retry(msgs, model=model, temperature=0.6, max_tokens=1500)
                ch_text += "\n\n" + dedup_boundary(ch_text, cont, lookback=180)
                passes += 1

            save_chapter_draft(manifest["title"], manifest["timestamp"], ch_num, ch_title, ch_text)
            manifest["completed"] = sorted(list(set(manifest.get("completed",[]) + [int(ch_num)])))
            manifest["next_index"] = idx + 1
            save_manifest(manifest["title"], manifest["timestamp"], manifest)

            manuscript_parts.append(f"\n\n# Chapter {ch_num}: {ch_title}\n\n{ch_text}\n")
            prior_text += f"\n\n[END CH {ch_num}]\n{ch_text}\n"

        progress.progress(1.0, text="Combining‚Ä¶")
        manuscript = "\n".join(manuscript_parts).strip()

        ts_final = int(time.time())
        safe_title = safe_title_slug(manifest["title"])
        md_name = f"{safe_title}_{ts_final}.md"
        docx_name = f"{safe_title}_{ts_final}.docx"

        st.download_button("‚¨áÔ∏è Download manuscript (Markdown)", manuscript, file_name=md_name)

        try:
            docx_bytes = md_to_docx(manuscript, manifest["title"], manifest.get("thesis",""), author=manifest.get("author",""))
            st.download_button("‚¨áÔ∏è Download manuscript (Word .docx)", docx_bytes, file_name=docx_name)
        except Exception as e:
            st.warning(f"Could not generate .docx: {e}")
            docx_bytes = None

        try:
            res = gh_put_file(f"outputs/{md_name}", manuscript.encode("utf-8"), f"Add manuscript {md_name}")
            owner, repo, branch = gh_repo_info()
            link_branch = branch or "main"
            gh_link_md = f"https://github.com/{owner}/{repo}/blob/{link_branch}/{res.get('content',{}).get('path', f'outputs/{md_name}')}"
            st.success(f"Saved Markdown to GitHub: {gh_link_md}")
        except Exception as e:
            st.warning(f"Could not save Markdown to GitHub automatically: {e}")

        try:
            if docx_bytes:
                res2 = gh_put_file(f"outputs/{docx_name}", docx_bytes, f"Add manuscript {docx_name}")
                owner, repo, branch = gh_repo_info()
                link_branch = branch or "main"
                gh_link_docx = f"https://github.com/{owner}/{repo}/blob/{link_branch}/{res2.get('content',{}).get('path', f'outputs/{docx_name}')}"
                st.success(f"Saved Word to GitHub: {gh_link_docx}")
        except Exception as e:
            st.warning(f"Could not save Word to GitHub automatically: {e}")

        st.success("‚úÖ Book generation complete.")
    except Exception as e:
        st.error(f"Generation stopped with error: {e}")
    finally:
        st.session_state.busy = False

# ========== Post-process helpers ==========
def _docx_to_md_simple(file_bytes: bytes) -> str:
    from docx import Document as _Doc
    doc = _Doc(io.BytesIO(file_bytes))
    lines = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            lines.append("")
            continue
        style = (p.style.name or "").lower()
        if "heading" in style:
            lvl = 1
            m = re.search(r"heading\s*(\d+)", style)
            if m:
                lvl = max(1, min(6, int(m.group(1))))
            lines.append("#" * lvl + " " + txt)
        else:
            lines.append(txt)
    return "\n".join(lines)

def _split_chapters(md_text: str):
    chunks = []
    cur_head = None
    cur = []
    for line in md_text.splitlines():
        if line.startswith("# Chapter "):
            if cur_head is not None:
                chunks.append((cur_head, "\n".join(cur).strip()))
                cur = []
            cur_head = line.strip()
        else:
            cur.append(line)
    if cur_head is not None:
        chunks.append((cur_head, "\n".join(cur).strip()))
    return chunks

def _ensure_single_conclusion(ch_md: str) -> str:
    parts = ch_md.split("\n")
    concl_idx = [i for i,l in enumerate(parts) if l.strip().lower() in ("## conclusion","## final conclusion")]
    if len(concl_idx) <= 1:
        return ch_md
    keep = []
    i = 0
    while i < len(parts):
        if i in concl_idx[:-1] and parts[i].strip().lower().startswith("##"):
            i += 1
            continue
        keep.append(parts[i]); i += 1
    return "\n".join(keep)

def _condense_subheads(ch_md: str, max_headings: int = 4) -> str:
    lines = ch_md.split("\n")
    idxs = [i for i,l in enumerate(lines) if l.strip().startswith("## ") and l.strip().lower() not in ("## conclusion","## final conclusion")]
    if len(idxs) <= max_headings:
        return ch_md
    keep_idxs = set(idxs[:max_headings])
    out = []
    for i, l in enumerate(lines):
        if i in idxs and i not in keep_idxs:
            continue
        out.append(l)
    return "\n".join(out)

def _dedup_paragraphs(ch_md: str) -> str:
    paras = [p.strip() for p in ch_md.split("\n\n")]
    seen = set(); out = []
    for p in paras:
        key = re.sub(r"\W+"," ", p.lower()).strip()[:280]
        if key and key not in seen:
            seen.add(key); out.append(p)
    return "\n\n".join(out)

def _normalize_heading_levels(ch_md: str) -> str:
    lines = ch_md.split("\n"); out = []
    for l in lines:
        if l.startswith("#####"): out.append("### " + l.lstrip("#").strip())
        elif l.startswith("####"): out.append("### " + l.lstrip("#").strip())
        elif l.startswith("### "): out.append(l)
        elif l.startswith("## "): out.append(l)
        elif l.startswith("# "): out.append("## " + l[2:])
        else: out.append(l)
    return "\n".join(out)

# ---- People & quotes balancing ----
def _extract_person_from_attrib_line(line: str) -> Optional[str]:
    m = re.search(r"‚Äî\s*([A-Z][A-Za-z]+(?:\s[A-Z][A-Za-z\-\']+)+)", line.strip())
    return m.group(1).strip() if m else None

def _quote_blocks(ch_md: str) -> List[Tuple[int,int,str,str]]:
    """
    Find quote blocks in chapter:
    returns list of (start_idx, end_idx_exclusive, quote_line, attrib_line)
    """
    lines = ch_md.split("\n")
    out = []
    i = 0
    while i < len(lines):
        if lines[i].strip().startswith("> ") and not lines[i].strip().startswith("> ‚Äî"):
            qline = lines[i]
            if i+1 < len(lines) and lines[i+1].strip().startswith("> ‚Äî"):
                out.append((i, i+2, qline, lines[i+1]))
                i += 2
                continue
        i += 1
    return out

def _build_people_index_from_records(k: int = 80) -> List[Tuple[str,str]]:
    if not records: return []
    hits = retrieve("leadership creativity service guest experience quotes exemplars hotels restaurants chefs", k=k, tags=["Books"])
    quotes = []
    used = set()
    for h in hits:
        txt = h["text"]; src = h.get("source","")
        spans = re.findall(r"[‚Äú\"]([^‚Äù\"]{40,420})[‚Äù\"]", txt)
        # heuristic: prefer true quotes; else fallback to a strong sentence
        if not spans:
            sents = re.split(r"(?<=[\.\!\?])\s+", txt.strip())
            sents = [s for s in sents if 60 <= len(s) <= 260]
            if sents: spans = [sents[0]]
        if spans and src not in used:
            q = "‚Äú" + spans[0].strip() + "‚Äù"
            # Try to infer person from source filename (e.g., 'Anneke Brown ‚Äî COMO The Treasury.docx')
            person = re.sub(r"\.\w+$","", src).strip()
            # Keep simple/clean names (at least two words)
            if len(person.split()) < 2:
                # attempt to infer from text (leading capitalized bigrams)
                m = re.search(r"\b([A-Z][a-z]+ [A-Z][A-Za-z\-\']+)", txt[:140])
                person = m.group(1) if m else "Source Book"
            quotes.append((q, f"‚Äî {person}"))
            used.add(src)
    return quotes

def _limit_overuse_and_topup_diversity(ch_head: str, ch_md: str, max_per_person: int, min_unique_people: int, max_inserts: int) -> str:
    """
    Cap repeated attributions to the same person and add new diverse quotes if chapter lacks variety.
    """
    lines = ch_md.split("\n")
    blocks = _quote_blocks(ch_md)
    counts: Dict[str,int] = {}
    # Count current people
    for _,_,_, attrib in blocks:
        who = _extract_person_from_attrib_line(attrib) or "Unknown"
        counts[who] = counts.get(who, 0) + 1
    # Remove excess quotes for overused people (keep earliest)
    to_remove_ranges = []
    seen_per_person: Dict[str,int] = {}
    for start, end, q, attrib in blocks:
        who = _extract_person_from_attrib_line(attrib) or "Unknown"
        seen_per_person[who] = seen_per_person.get(who, 0) + 1
        if max_per_person > 0 and seen_per_person[who] > max_per_person:
            to_remove_ranges.append((start, end))
    # Apply removals (back to front)
    for s,e in sorted(to_remove_ranges, key=lambda x: -x[0]):
        del lines[s:e]
    ch_md = "\n".join(lines)
    # Recalculate after removals
    blocks = _quote_blocks(ch_md)
    people_now = set()
    for _,_,_, attrib in blocks:
        who = _extract_person_from_attrib_line(attrib) or "Unknown"
        people_now.add(who)
    # Top-up with diverse quotes if needed
    need = max(0, min_unique_people - len([p for p in people_now if p != "Unknown"]))
    if need > 0 and max_inserts > 0:
        candidates = _build_people_index_from_records(k=120)
        # filter out already present people
        candidates = [(q,a) for (q,a) in candidates if (_extract_person_from_attrib_line(a) or "") not in people_now]
        inserts = candidates[:min(need, max_inserts)]
        if inserts:
            ins_lines = ch_md.split("\n")
            # insert after first H2 section start
            insert_at = 0
            for i,l in enumerate(ins_lines):
                if l.strip().startswith("## ") and l.strip().lower() not in ("## conclusion","## final conclusion"):
                    insert_at = i+1; break
            block = []
            for q, a in inserts:
                block += ["> " + q, "> " + a, ""]
            ins_lines = ins_lines[:insert_at] + [""] + block + ins_lines[insert_at:]
            ch_md = "\n".join(ins_lines)
    return ch_md

def _insert_quotes_into_chapter(ch_head: str, ch_md: str, min_q: int, max_q: int) -> str:
    if max_q <= 0 or min_q <= 0: return ch_md
    title = re.sub(r"^#*\s*Chapter\s+\d+:\s*","", ch_head).strip()
    want = max(min_q, 1)
    # pull quotes favoring this chapter's theme via retrieval
    if records:
        hits = retrieve(f'{title} leadership creativity service guest experience examples quotes', k=18, tags=["Books"])
    else:
        hits = []
    quotes = []; used_sources = set()
    for h in hits:
        txt = h["text"]; src = h.get("source","")
        spans = re.findall(r"[‚Äú\"]([^‚Äù\"]{40,420})[‚Äù\"]", txt)
        person = re.sub(r"\.\w+$","", src).strip() if src else ""
        if not spans:
            sents = re.split(r"(?<=[\.\!\?])\s+", txt.strip())
            sents = [s for s in sents if 60 <= len(s) <= 260]
            if sents: spans = [sents[0]]
        if spans and (src not in used_sources):
            q = spans[0].strip()
            attrib = person or (src if src else "Source Book")
            line = f"‚Äî {attrib}"
            quotes.append((f"‚Äú{q}‚Äù", line))
            used_sources.add(src)
        if len(quotes) >= want: break
    # fallback: nothing found
    if not quotes:
        return ch_md
    lines = ch_md.split("\n")
    insert_at = 0
    for i, l in enumerate(lines):
        if l.strip().startswith("## ") and l.strip().lower() not in ("## conclusion","## final conclusion"):
            insert_at = i + 1; break
    block = []
    for q, who in quotes[:max_q]:
        block.append("> " + q); block.append("> " + who); block.append("")
    new_lines = lines[:insert_at] + [""] + block + lines[insert_at:]
    return "\n".join(new_lines)

def _rebuild_docx_and_save(title: str, thesis_local: str, author: str, md_text: str, out_prefix: str) -> Tuple[Optional[str], Optional[str], Optional[str]]:
    ts_final = int(time.time())
    safe_title = safe_title_slug(title)
    md_name = f"{safe_title}_{out_prefix}_{ts_final}.md"
    docx_name = f"{safe_title}_{out_prefix}_{ts_final}.docx"
    try:
        res_md = gh_put_file(f"outputs/{md_name}", md_text.encode("utf-8"), f"Add postprocessed {md_name}")
        owner, repo, branch = gh_repo_info(); link_branch = branch or "main"
        md_link = f"https://github.com/{owner}/{repo}/blob/{link_branch}/{res_md.get('content',{}).get('path', f'outputs/{md_name}')}"
    except Exception:
        md_link = None
    try:
        docx_bytes = md_to_docx(md_text, title, thesis_local, author=author)
        res_docx = gh_put_file(f"outputs/{docx_name}", docx_bytes, f"Add postprocessed {docx_name}")
        owner, repo, branch = gh_repo_info(); link_branch = branch or "main"
        docx_link = f"https://github.com/{owner}/{repo}/blob/{link_branch}/{res_docx.get('content',{}).get('path', f'outputs/{docx_name}')}"
    except Exception:
        docx_link = None
    return md_link, docx_link, None

# ---- Flow polish + length enforcement ----
def _polish_flow(ch_head: str, ch_md: str, target_words: int, max_headings: int, model: str) -> str:
    prompt = f"""
You are a senior line editor. Rewrite the following chapter text to:
- streamline flow and transitions, remove clunky phrasing and redundancies,
- keep a SINGLE "## Conclusion" at the end,
- limit H2 sections (##) to ‚â§ {max_headings} total (excluding Conclusion),
- preserve existing hyperlinks and all verbatim quote attributions of the form: ‚Äú...‚Äù ‚Äî Full Name, Role (Country),
- maintain the author‚Äôs voice and factual content,
- aim for ~{target_words} words.

Return ONLY the revised Markdown content for this chapter body (do not add or remove chapter headings outside its body).
Text:
{ch_md}
""".strip()
    msgs = [
        {"role":"system","content":"You are a precise nonfiction editor. Reply in Markdown only."},
        {"role":"user","content":prompt}
    ]
    return chat_with_retry(msgs, model=model, temperature=0.3, max_tokens=2800)

def _enforce_length(ch_md: str, target_words: int, tolerance_pct: int, model: str) -> str:
    tgt = int(target_words)
    tol = max(0, min(80, int(tolerance_pct)))
    lo = int(tgt * (1 - tol/100))
    hi = int(tgt * (1 + tol/100))
    count = wc(ch_md)
    if lo <= count <= hi:
        return ch_md
    if count > hi:
        prompt = f"""
Condense the following chapter to ~{tgt} words (¬±{tol}%). Remove repetition and tighten language.
Keep the same headings, hyperlinks, and all quote attributions (‚Äú‚Ä¶‚Äù ‚Äî Name).
Return only the revised Markdown body.

{ch_md}
""".strip()
        msgs = [
            {"role":"system","content":"You are a skilled compressor of nonfiction prose. Reply in Markdown only."},
            {"role":"user","content":prompt}
        ]
        return chat_with_retry(msgs, model=model, temperature=0.3, max_tokens=2600)
    else:
        prompt = f"""
Expand the following chapter to ~{tgt} words (¬±{tol}%) by deepening analysis and examples (paraphrase from existing material if needed).
Do NOT invent new quotes or links. Preserve all existing headings, hyperlinks, and quote attributions.
Return only the revised Markdown body.

{ch_md}
""".strip()
        msgs = [
            {"role":"system","content":"You expand thoughtfully without fluff. Reply in Markdown only."},
            {"role":"user","content":prompt}
        ]
        return chat_with_retry(msgs, model=model, temperature=0.4, max_tokens=2800)

# ========== UI TABS ==========
tab_gen, tab_post = st.tabs(["üõ† Generate", "üßπ Post‚Äëprocess"])

with tab_gen:
    st.subheader("1) (Optional) Upload source books or materials")
    files = st.file_uploader("PDF/DOCX/MD/TXT ‚Äî your books or any key sources",
                             type=["pdf","docx","md","markdown","txt"], accept_multiple_files=True, disabled=st.session_state.busy)
    tags_str = st.text_input("Tags for these files (comma-separated)", value="Books", disabled=st.session_state.busy)
    persist = st.checkbox("Save originals to GitHub (/uploads)", value=True, disabled=st.session_state.busy)

    if st.button("Ingest selected files", disabled=st.session_state.busy):
        if not files:
            st.warning("No files selected.")
        else:
            tags = [t.strip() for t in tags_str.split(",") if t.strip()] or ["Books"]
            added = 0
            for f in files:
                b = f.read()
                tmp = os.path.join("data", f"{int(time.time())}_{f.name}")
                with open(tmp, "wb") as out:
                    out.write(b)
                if persist:
                    try:
                        gh_put_file(f"uploads/{int(time.time())}_{f.name}", b, f"Add source {f.name}")
                    except Exception as e:
                        st.error(f"GitHub upload failed for {f.name}: {e}")
                try:
                    raw = load_text_from_file(tmp)
                    pieces = chunk_text(raw, chunk_size=1000, overlap=150)
                    embs = embed_texts(pieces, model=st.session_state.emb_model)
                    for i,(txt,emb) in enumerate(zip(pieces,embs)):
                        records.append({
                            "id": f"{sha16(f.name)}:{i}:{len(records)}",
                            "text": txt, "tags": tags, "source": f.name, "emb": emb,
                        })
                    added += len(embs)
                except Exception as e:
                    st.error(f"Failed to process {f.name}: {e}")
            st.success(f"Ingested: {len(files)} file(s) ‚Üí {added} chunks.")
    st.caption(f"Corpus size: {len(records)} chunks. (Zero is okay ‚Äî model + web URLs only mode)")

    st.subheader("2) Define the project")
    col1, col2 = st.columns(2)
    with col1:
        book_title = st.text_input("Book title", value="Working Title: The Essence of Modern Luxury", disabled=st.session_state.busy)
        thesis = st.text_area("Book thesis / goal (one paragraph)", height=120, disabled=st.session_state.busy)
    with col2:
        style = st.text_area("Voice & Style Sheet (tense, pacing, terminology, audience)", height=120, disabled=st.session_state.busy)

    st.subheader("3) Outline & generation settings")
    outline_mode = st.radio("Outline source", ["Auto-generate", "Paste outline", "Upload outline file"],
                            index=0, horizontal=True, disabled=st.session_state.busy)
    num_chapters = st.number_input("Total number of chapters (Intro is #1, Final Conclusion is last)",
                                   8, 24, 12, 1, disabled=st.session_state.busy)
    words_per_chapter = st.number_input("Target words per chapter", 1200, 10000, 3500, 100, disabled=st.session_state.busy)

    st.markdown("**Web research** (optional)")
    use_web = st.checkbox("Enable web research with manual URLs (below)", disabled=st.session_state.busy)
    ref_urls = st.text_area("Manual Reference URLs (one per line)",
                            placeholder="https://example.com/report\nhttps://another.com/profile",
                            height=100, disabled=st.session_state.busy)
    urls_list = [u.strip() for u in ref_urls.splitlines() if u.strip()] if use_web else []

    st.markdown("**Web research guidance (Tavily)**")
    research_guidance = st.text_area(
        "Tell the system how to search and what to prioritize (e.g., peer-reviewed studies since 2019; include hyperlinks to journals/DOIs; favor meta-analyses).",
        height=120, placeholder="Prioritize recent peer-reviewed research (2019+), systematic reviews, and meta-analyses. Add inline links to journals/DOIs. Include major industry reports if relevant.",
        disabled=st.session_state.busy
    )
    use_tavily = st.checkbox("Use Tavily web research automatically", value=True, disabled=st.session_state.busy)
    max_sources_per_chapter = st.slider("Max Tavily sources per chapter", 2, 20, 6, 1, disabled=st.session_state.busy)
    recent_year_cutoff = st.number_input("Prefer sources published ‚â• this year (0 to ignore)", min_value=0, max_value=2100, value=2019, step=1, disabled=st.session_state.busy)
    include_domains_txt = st.text_input("Include only these domains (comma-separated, optional)", value="", disabled=st.session_state.busy)
    exclude_domains_txt = st.text_input("Exclude these domains (comma-separated, optional)", value="", disabled=st.session_state.busy)

    st.divider()
    colA, colB = st.columns([1,1])
    with colA:
        generate = st.button("üöÄ One-Click: Generate Entire Book", type="primary", disabled=st.session_state.busy, use_container_width=True)
    with colB:
        resume = st.button("‚ñ∂Ô∏è Resume last run (from GitHub drafts)", disabled=st.session_state.busy, use_container_width=True)

    if generate:
        generate_book(
            run_mode="fresh",
            model=st.session_state.model,
            emb_model=st.session_state.emb_model,
            records=records,
            book_title=book_title, thesis=thesis, style=style,
            outline_mode=outline_mode, pasted_outline=(st.text_area("Paste outline (only used if 'Paste outline' is selected)", height=120) if outline_mode=="Paste outline" else ""),
            outline_file=(st.file_uploader("Upload outline file (only used if 'Upload outline file' is selected)", type=["txt","md","markdown","docx","pdf"]) if outline_mode=="Upload outline file" else None),
            num_chapters=int(num_chapters), words_per_chapter=int(words_per_chapter),
            top_k=10, chunk_size=1000, chunk_overlap=150,
            use_tavily=use_tavily, research_guidance=research_guidance, recent_year_cutoff=int(recent_year_cutoff),
            include_domains_txt=include_domains_txt, exclude_domains_txt=exclude_domains_txt,
            urls_list=urls_list, max_sources_per_chapter=int(max_sources_per_chapter),
            temperature=0.6
        )

    if resume:
        generate_book(
            run_mode="resume",
            model=st.session_state.model,
            emb_model=st.session_state.emb_model,
            records=records,
            book_title="", thesis="", style="",
            outline_mode="", pasted_outline="", outline_file=None,
            num_chapters=0, words_per_chapter=0,
            top_k=10, chunk_size=1000, chunk_overlap=150,
            use_tavily=True, research_guidance="", recent_year_cutoff=0,
            include_domains_txt="", exclude_domains_txt="",
            urls_list=[], max_sources_per_chapter=6,
            temperature=0.6
        )

with tab_post:
    st.subheader("4) Post‚Äëprocess manuscript (cleanup + quotes + diversity + flow + length)")
    st.markdown("Upload the completed draft (**.docx**, **.md**, or **.txt**). Choose clean‚Äëups; we‚Äôll diversify quotes, streamline flow, enforce chapter length targets, and output a revised Word/Markdown + diff. Quotes are pulled from your ingested books when available.")

    pp_file = st.file_uploader("Upload manuscript (.docx / .md / .txt)", type=["docx","md","txt"], accept_multiple_files=False, key="pp_upl")

    st.markdown("**Cleanup options**")
    col_pp1, col_pp2, col_pp3 = st.columns(3)
    with col_pp1:
        fix_conclusion = st.checkbox("Ensure single Conclusion per chapter", value=True)
        condense_heads = st.checkbox("Condense H2 subheadings (‚â§ N)", value=True)
        max_h2 = st.number_input("Max H2 per chapter (excl. Conclusion)", 2, 12, 4, 1)
    with col_pp2:
        dedup_redund = st.checkbox("Remove redundant paragraphs", value=True)
        normalize_heads = st.checkbox("Normalize heading levels", value=True)
        quote_enable = st.checkbox("Insert missing quotes from books", value=True)
    with col_pp3:
        cap_overuse = st.checkbox("Cap overuse of the same person", value=True)
        max_per_person = st.number_input("Max quotes per person / chapter", 1, 6, 2, 1)
        min_unique_people = st.number_input("Min unique people / chapter", 0, 10, 2, 1)

    st.markdown("**Polish & length**")
    col_len1, col_len2 = st.columns(2)
    with col_len1:
        enforce_length = st.checkbox("Enforce chapter length target", value=True)
        target_words = st.number_input("Target words / chapter", 800, 12000, 3500, 50)
        tolerance_pct = st.slider("Tolerance ¬±%", 0, 50, 10, 1)
    with col_len2:
        polish_flow = st.checkbox("Polish flow & transitions (LLM)", value=True)
        max_inserts = st.number_input("Max diversity inserts per chapter", 0, 10, 2, 1)

    st.caption("Tip: Keep diversity caps on to avoid over-relying on a single interviewee. Length enforcement respects your target within the chosen tolerance.")
    pp_go = st.button("üßπ Clean, Balance & Polish", type="primary", disabled=st.session_state.busy)

    if pp_go:
        if not pp_file:
            st.warning("Please upload a manuscript file.")
        else:
            try:
                # Load manuscript into Markdown
                name = pp_file.name.lower()
                if name.endswith(".docx"):
                    md_raw = _docx_to_md_simple(pp_file.read())
                elif name.endswith(".md"):
                    md_raw = pp_file.read().decode("utf-8", errors="ignore")
                else:  # .txt
                    md_raw = pp_file.read().decode("utf-8", errors="ignore")

                md_orig = md_raw

                header, chapters = "", []
                if "\n# Chapter " in md_orig:
                    header, rest = md_orig.split("\n# Chapter ", 1)
                    chapters = _split_chapters("# Chapter " + rest)
                else:
                    chapters = [("# Chapter 1: Untitled", md_orig)]
                    header = ""

                revised = []
                import difflib as _dif

                for head, body in chapters:
                    ch = body
                    if normalize_heads:
                        ch = _normalize_heading_levels(ch)
                    if fix_conclusion:
                        ch = _ensure_single_conclusion(ch)
                    if condense_heads:
                        ch = _condense_subheads(ch, max_headings=int(max_h2))
                    if dedup_redund:
                        ch = _dedup_paragraphs(ch)
                    if cap_overuse:
                        ch = _limit_overuse_and_topup_diversity(head, ch,
                                                                max_per_person=int(max_per_person),
                                                                min_unique_people=int(min_unique_people),
                                                                max_inserts=int(max_inserts))
                    if quote_enable:
                        # add quotes if chapter has none
                        if not _quote_blocks(ch):
                            ch = _insert_quotes_into_chapter(head, ch, min_q=1, max_q=2)

                    if polish_flow:
                        ch = _polish_flow(head, ch, target_words=int(target_words), max_headings=int(max_h2), model=st.session_state.model)

                    if enforce_length:
                        ch = _enforce_length(ch, target_words=int(target_words), tolerance_pct=int(tolerance_pct), model=st.session_state.model)

                    # Always ensure exactly one Conclusion at end post‚Äëedits
                    ch = _ensure_single_conclusion(ch)
                    # Clip subheads again if polish/length added extra
                    ch = _condense_subheads(ch, max_headings=int(max_h2))

                    revised.append((head, ch))

                manuscript_pp = (header.strip() + "\n\n" if header.strip() else "") + "\n\n".join(
                    [f"{h}\n\n{b}".strip() for (h,b) in revised]
                ).strip()

                diff_text = "\n".join(_dif.unified_diff(md_orig.splitlines(), manuscript_pp.splitlines(), lineterm="", fromfile="original.md", tofile="revised.md"))

                st.download_button("‚¨áÔ∏è Download cleaned manuscript (Markdown)", manuscript_pp, file_name="manuscript_cleaned.md")
                try:
                    docx_bytes_pp = md_to_docx(manuscript_pp, "Revised Manuscript", "", author=_sec("BOOK_AUTHOR") or "")
                    st.download_button("‚¨áÔ∏è Download cleaned manuscript (Word .docx)", docx_bytes_pp, file_name="manuscript_cleaned.docx")
                except Exception as e:
                    st.warning(f"Could not make .docx: {e}")
                st.download_button("‚¨áÔ∏è Download diff (MD)", diff_text, file_name="manuscript_diff.md")

                md_link, docx_link, _ = _rebuild_docx_and_save("Revised Manuscript", "", _sec("BOOK_AUTHOR") or "", manuscript_pp, out_prefix="POST")
                if md_link: st.success(f"Saved cleaned Markdown to GitHub: {md_link}")
                if docx_link: st.success(f"Saved cleaned Word to GitHub: {docx_link}")

                st.success("‚úÖ Post‚Äëprocessing complete.")
            except Exception as e:
                st.error(f"Post‚Äëprocessing failed: {e}")
